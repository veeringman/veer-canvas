"""Extract plain text from uploaded / Drive files for the composer."""

from __future__ import annotations

import io
import pathlib
import re
import zipfile
from html import unescape
from typing import Any

IMPORT_MAX_BYTES = 12 * 1024 * 1024
IMPORT_EXTS = {".txt", ".text", ".doc", ".docx", ".pages", ".pdf"}
_PAGES_PREVIEW = (
    "preview.pdf",
    "Preview.pdf",
    "QuickLook/Preview.pdf",
    "preview/preview.pdf",
)


def _html_escape(text: str) -> str:
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def text_to_html(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not raw:
        return "<p></p>"
    blocks = re.split(r"\n\s*\n", raw)
    parts: list[str] = []
    for block in blocks:
        lines = [_html_escape(line) for line in block.split("\n")]
        parts.append("<p>" + "<br>".join(lines) + "</p>")
    return "".join(parts) or "<p></p>"


def decode_text_bytes(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_pdf_text(data: bytes, *, max_chars: int = 120000) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise ValueError("PDF text extract needs pypdf on the server.") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValueError("Could not read that PDF.") from exc
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        try:
            chunk = (page.extract_text() or "").strip()
        except Exception:
            chunk = ""
        if not chunk:
            continue
        parts.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("No extractable text in that PDF (it may be a scan).")
    return text[:max_chars]


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise ValueError("Word import needs python-docx on the server.") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError("Could not read that Word file. Try saving as .docx.") from exc
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").replace("\n", " ").strip()) for cell in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("That Word file has no extractable text.")
    return text


def extract_html_text(raw: str) -> str:
    html = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw or "")
    html = re.sub(r"(?i)<br\s*/?>", "\n", html)
    html = re.sub(r"(?i)</(p|div|h[1-6]|li|tr)>", "\n", html)
    html = re.sub(r"<[^>]+>", " ", html)
    html = unescape(html).replace("\xa0", " ")
    html = re.sub(r"[ \t]+", " ", html)
    html = re.sub(r"\n{3,}", "\n\n", html)
    return html.strip()


def extract_doc_text(data: bytes) -> str:
    if data[:2] == b"PK":
        return extract_docx_text(data)
    head = data[:8000].lower()
    if b"<html" in head or b"urn:schemas-microsoft-com:office:word" in head:
        return extract_html_text(decode_text_bytes(data))
    raise ValueError("Older .doc files need to be saved as .docx (or PDF) first.")


def extract_pages_text(data: bytes) -> str:
    bio = io.BytesIO(data)
    if not zipfile.is_zipfile(bio):
        raise ValueError("Could not read that Pages file. Export as PDF or Word.")
    bio.seek(0)
    with zipfile.ZipFile(bio) as zf:
        names = zf.namelist()
        total = 0
        for info in zf.infolist():
            total += int(info.file_size or 0)
            if total > 40 * 1024 * 1024:
                raise ValueError("That Pages file is too large to import.")
        for cand in _PAGES_PREVIEW:
            if cand in names:
                return extract_pdf_text(zf.read(cand))
        xmls = [n for n in names if n.lower().endswith("index.xml")]
        if xmls:
            raw = zf.read(xmls[0])
            text = extract_html_text(decode_text_bytes(raw))
            if text:
                return text
    raise ValueError("No extractable text in that Pages file. Export as PDF or Word.")


def sniff_kind(filename: str, mime: str | None, data: bytes) -> str:
    name = (filename or "").lower()
    suffix = pathlib.Path(name).suffix.lower()
    mime_l = (mime or "").lower()
    if suffix in IMPORT_EXTS:
        return suffix.lstrip(".")
    if "pdf" in mime_l:
        return "pdf"
    if "wordprocessingml" in mime_l or suffix == ".docx":
        return "docx"
    if mime_l == "application/msword" or suffix == ".doc":
        return "doc"
    if "pages" in mime_l or suffix == ".pages":
        return "pages"
    if mime_l.startswith("text/") or suffix in {".txt", ".text"}:
        return "txt"
    if "google-apps.document" in mime_l:
        return "gdoc"
    if data[:5] == b"%PDF-":
        return "pdf"
    if data[:2] == b"PK":
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = zf.namelist()
            if any(n in names for n in _PAGES_PREVIEW) or any(n.lower().endswith("index.xml") for n in names):
                return "pages"
        except Exception:
            pass
        return "docx"
    return suffix.lstrip(".") or "bin"


def extract_import(
    data: bytes,
    *,
    filename: str = "",
    mime: str | None = None,
) -> dict[str, Any]:
    if not data:
        raise ValueError("Empty file")
    if len(data) > IMPORT_MAX_BYTES:
        raise ValueError("File too large to import (max 12 MB)")
    kind = sniff_kind(filename, mime, data)
    if kind == "gdoc":
        text = decode_text_bytes(data).strip()
    elif kind in {"txt", "text"}:
        text = decode_text_bytes(data).strip()
    elif kind == "pdf":
        text = extract_pdf_text(data)
    elif kind == "docx":
        text = extract_docx_text(data)
    elif kind == "doc":
        text = extract_doc_text(data)
    elif kind == "pages":
        text = extract_pages_text(data)
    else:
        raise ValueError("Use a .txt, Word (.doc/.docx), Pages, or PDF file.")
    text = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
    if not text:
        raise ValueError("No extractable text in that file.")
    stem = pathlib.Path(filename or "Imported").stem or "Imported"
    return {
        "text": text,
        "html": text_to_html(text),
        "title": stem[:160],
        "sourceName": pathlib.Path(filename or "document").name,
        "kind": kind,
    }
