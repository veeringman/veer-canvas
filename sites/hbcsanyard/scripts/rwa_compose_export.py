"""Compose export helpers — HTML fragment to text / Word, pad body injection."""

from __future__ import annotations

import base64
import io
import pathlib
import re
from html import unescape
from html.parser import HTMLParser
from typing import Any

DIAMOND_RULE_HTML = '<div class="rule" aria-hidden="true"><span class="pip"></span></div>'


def _html_escape(text: Any) -> str:
    s = "" if text is None else str(text)
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


_CSS_DECL_RE = re.compile(r"([a-z-]+)\s*:\s*([^;]+)", re.I)


def _css_decls(style: str) -> dict[str, str]:
    return {m.group(1).strip().lower(): m.group(2).strip() for m in _CSS_DECL_RE.finditer(style or "")}


def _css_hex(value: str) -> str | None:
    s = (value or "").strip().lower()
    if s in {"none", "transparent"}:
        return None
    if re.fullmatch(r"#[0-9a-f]{6}", s):
        return s[1:]
    if re.fullmatch(r"#[0-9a-f]{3}", s):
        return "".join(ch * 2 for ch in s[1:])
    m = re.match(r"rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", s)
    if not m:
        return None
    if s.startswith("rgba") and re.search(r",\s*0(?:\.0+)?\s*\)$", s):
        return None
    return "".join(f"{int(m.group(i)):02x}" for i in range(1, 4))


def _css_pt(value: str) -> float | None:
    s = (value or "").strip().lower()
    if s in {"", "none"}:
        return 0.0 if s == "none" else None
    m = re.match(r"([0-9.]+)\s*(pt|px)?", s)
    if not m:
        return None
    n = float(m.group(1))
    if m.group(2) == "px":
        n = n * 72.0 / 96.0
    return n


def _apply_docx_cell(cell: Any, spec: dict[str, str]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    text = spec.get("text") or ""
    cell.text = text
    decls = _css_decls(spec.get("style") or "")
    fill = _css_hex(decls.get("background-color") or decls.get("background") or "")
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    border = decls.get("border") or ""
    bwidth = _css_pt(decls.get("border-width") or "")
    if border.strip() == "none" or "border-style: none" in (spec.get("style") or "").lower():
        bwidth = 0.0
    if bwidth is None:
        for part in border.split():
            got = _css_pt(part)
            if got is not None:
                bwidth = got
                break
    bcolor = _css_hex(decls.get("border-color") or "")
    if not bcolor:
        for part in border.split():
            got = _css_hex(part)
            if got:
                bcolor = got
                break
    if bwidth is None and not border:
        bwidth = 0.6
        bcolor = bcolor or "0b2a56"
    if bwidth is not None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        val = "nil" if bwidth <= 0 else "single"
        sz = str(max(0, int(round(bwidth * 8))))
        for edge in ("top", "left", "bottom", "right"):
            edge_el = OxmlElement(f"w:{edge}")
            edge_el.set(qn("w:val"), val)
            if val != "nil":
                edge_el.set(qn("w:sz"), sz)
                edge_el.set(qn("w:space"), "0")
                edge_el.set(qn("w:color"), bcolor or "0b2a56")
            borders.append(edge_el)
        tc_pr.append(borders)
    color = _css_hex(decls.get("color") or "")
    if color:
        try:
            rgb = RGBColor.from_string(color)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = rgb
        except Exception:
            pass

COMPOSE_PAD_BODY_CSS = """
<style id="mhws-compose-body">
  .screen-hint, .layout-picker { display: none !important; }
  .body-area p { margin: 0 0 8pt; }
  .body-area h2 {
    margin: 0 0 8pt;
    font-size: 18pt;
    font-weight: 700;
    line-height: 1.25;
    color: #0b2a56;
  }
  .body-area h3 {
    margin: 0 0 6pt;
    font-size: 14pt;
    font-weight: 700;
    line-height: 1.3;
    color: #143a6e;
  }
  .body-area h2 span,
  .body-area h3 span { font-size: inherit; font-weight: inherit; color: inherit; }
  .body-area .mhws-tab { white-space: pre; tab-size: 4; }
  .body-area .mhws-img-pair {
    display: flex;
    align-items: stretch;
    gap: 10pt;
    width: 100%;
    margin: 0 0 8pt;
  }
  .body-area .mhws-img-text { flex: 1 1 auto; min-width: 0; }
  .body-area table { border-collapse: collapse; width: 100%; margin: 8pt 0; }
  .body-area th, .body-area td { border: 0.6pt solid #0b2a56; padding: 4pt 6pt; vertical-align: top; }
  .body-area th { background: #eef2f8; }
  .body-area table.mhws-table-noborder th,
  .body-area table.mhws-table-noborder td { border: 0 !important; }
  .body-area img { max-width: 100%; height: auto; }
  .body-area .mhws-img { max-width: 100%; }
  .body-area .mhws-img img { width: 100%; height: auto; display: block; }
  .brand .logo, header.org img, .org img {
    display: block;
    width: 24mm;
    height: auto;
    visibility: visible;
    border: 0;
    outline: 0;
    box-shadow: none;
    background: transparent;
  }
  img.wm { display: block; }
  @media print {
    body { background: #fff !important; }
    .sheet { margin: 0 !important; border: 0 !important; box-shadow: none !important; }
  }
</style>
"""

COMPOSE_CHROME_LAYOUT_CSS = """
<style id="mhws-compose-chrome-layout">
  .sheet[data-layout="top"] .officers,
  .officers {
    display: grid !important;
    grid-template-columns: repeat(4, 1fr);
    column-gap: 0;
    align-items: start;
    justify-items: center;
    margin: 0 0 1.6mm;
    width: 100%;
    padding: 0;
  }
  .sheet[data-layout="top"] .role,
  .officers .role {
    text-align: center;
    min-width: 0;
    width: 100%;
    padding: 0 2mm;
    position: relative;
    align-self: start;
  }
  .sheet[data-layout="top"] .role:not(:last-child)::after,
  .officers .role:not(:last-child)::after {
    content: "";
    position: absolute;
    top: 10%;
    bottom: 10%;
    right: 0;
    width: 0;
    border-right: 0.7pt solid rgba(11, 42, 86, 0.14);
  }
  .sheet[data-layout="top"] .officers + .rule,
  .officers + .rule {
    display: grid !important;
    margin: 0 0 1.6mm;
  }
  .officers-foot,
  .mhws-header-gold-rule {
    display: none !important;
  }
  .rule {
    display: grid !important;
    grid-template-columns: 1fr auto 1fr;
    align-items: center;
    gap: 3mm;
    margin: 0 0 2.2mm;
    width: 100%;
  }
  .rule::before,
  .rule::after {
    content: "";
    height: 0;
    border-top: 1pt solid #0b2a56;
  }
  .rule .pip {
    width: 2.2mm;
    height: 2.2mm;
    background: #c9a227;
    transform: rotate(45deg);
    box-shadow: 0 0 0 1.2pt #fff, 0 0 0 1.7pt rgba(11, 42, 86, 0.35);
  }
  img.wm, .wm {
    top: 50% !important;
  }
  .mhws-run-header .head,
  .mhws-print-head .head,
  .mhws-page-chrome-tpl .head,
  .mhws-chrome-inner .head,
  .mhws-run-header .mhws-st-head,
  .mhws-print-head .mhws-st-head,
  .mhws-page-chrome-tpl .mhws-st-head,
  .mhws-chrome-inner .mhws-st-head,
  .mhws-run-header .mhws-simple-head,
  .mhws-print-head .mhws-simple-head,
  .mhws-page-chrome-tpl .mhws-simple-head,
  .mhws-chrome-inner .mhws-simple-head,
  .mhws-run-header .org,
  .mhws-print-head .org,
  .mhws-page-chrome-tpl .org,
  .mhws-chrome-inner .org,
  .mhws-run-header .brand,
  .mhws-print-head .brand,
  .mhws-page-chrome-tpl .brand,
  .mhws-chrome-inner .brand {
    border-bottom: none !important;
  }
  .mhws-run-header .org,
  .mhws-print-head .org {
    padding: 5mm 12mm 2pt !important;
  }
  .mhws-run-header .org img,
  .mhws-print-head .org img {
    width: 14mm !important;
  }
  .mhws-run-header .org h1,
  .mhws-print-head .org h1 {
    margin: 3pt 0 0 !important;
    font-size: 13pt !important;
  }
  .mhws-run-header .org .sub,
  .mhws-print-head .org .sub {
    margin: 1.5pt 0 0 !important;
    font-size: 9pt !important;
  }
  .mhws-run-header .org .meta,
  .mhws-print-head .org .meta {
    margin: 1.5pt 0 0 !important;
    font-size: 7.5pt !important;
  }
  .mhws-run-header .mhws-simple-head,
  .mhws-print-head .mhws-simple-head {
    padding: 3mm 12mm 1.5mm !important;
  }
  .mhws-run-header .mhws-simple-head img,
  .mhws-print-head .mhws-simple-head img {
    width: 14mm !important;
  }
  .mhws-run-header .mhws-simple-head h1,
  .mhws-print-head .mhws-simple-head h1 {
    margin: 1.5mm 0 0 !important;
    font-size: 11pt !important;
  }
  .mhws-run-header .mhws-simple-chrome,
  .mhws-print-head .mhws-simple-chrome {
    box-sizing: border-box;
    width: 100%;
  }
  .mhws-run-header .mhws-simple-chrome .rule,
  .mhws-print-head .mhws-simple-chrome .rule {
    margin: 0 12mm 1.6mm !important;
    width: auto !important;
  }
</style>
"""

COMPOSE_PDF_CSS = """
<style id="mhws-compose-pdf">
  @page { size: 210mm 297mm; margin: 0; }
  html, body,
  html.pad-a4-full, html.pad-a4-full body,
  html.pad-a4-blank, html.pad-a4-blank body,
  html.mhws-compose-multipage, html.mhws-compose-multipage body {
    background: #fff !important;
    width: 210mm !important;
    min-height: 0 !important;
    height: auto !important;
    max-height: none !important;
    margin: 0 !important;
    overflow: visible !important;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
  }
  .screen-hint, .layout-picker { display: none !important; }
  .sheet,
  html.pad-a4-full .sheet,
  html.pad-a4-blank .sheet {
    position: relative !important;
    box-sizing: border-box !important;
    width: 210mm !important;
    min-height: 0 !important;
    height: auto !important;
    max-height: none !important;
    margin: 0 !important;
    border: 0 !important;
    overflow: visible !important;
    box-shadow: none !important;
    display: block !important;
  }
  .pad { display: block !important; flex: none !important; min-height: 0 !important; height: auto !important; overflow: visible !important; }
  .body-area, .body {
    flex: none !important;
    min-height: 0 !important;
    max-height: none !important;
    height: auto !important;
    overflow: visible !important;
  }
  html.is-mhws-paged .sheet {
    display: none !important;
    height: 0 !important;
    min-height: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: hidden !important;
  }
  html.is-mhws-paged .mhws-print-desk {
    display: block !important;
    width: 100% !important;
    box-sizing: border-box;
    padding: 8mm 0 !important;
    background: #c5cdd8 !important;
  }
  html.is-mhws-paged,
  html.is-mhws-paged body {
    width: auto !important;
    min-width: 100% !important;
    max-width: none !important;
    background: #c5cdd8 !important;
  }
  html.is-mhws-paged .mhws-print-sheet {
    width: 210mm !important;
    min-height: 297mm !important;
    height: 297mm !important;
    max-height: none !important;
    margin: 0 auto 8mm !important;
    box-shadow: 0 1px 6px rgba(15, 40, 80, 0.2) !important;
    page-break-after: avoid !important;
    break-after: avoid-page !important;
    overflow: visible !important;
  }
  html.is-mhws-paged .mhws-print-head,
  html.is-mhws-paged .mhws-print-foot {
    overflow: visible !important;
    flex-shrink: 0 !important;
  }
  html.is-mhws-paged .mhws-print-head .sheet {
    border: none !important;
    box-shadow: none !important;
  }
  html.is-mhws-paged .mhws-print-body {
    overflow: hidden !important;
    flex: 1 1 auto !important;
    min-height: 0 !important;
  }
  html.is-mhws-paged .mhws-print-sheet > img.wm {
    display: block !important;
    visibility: visible !important;
    opacity: 0.75 !important;
    position: absolute !important;
    left: 50% !important;
    top: 50% !important;
    transform: translate(-50%, -50%) !important;
    width: min(112mm, 70%) !important;
    max-height: 46% !important;
    height: auto !important;
    object-fit: contain !important;
    z-index: 0 !important;
    pointer-events: none !important;
  }
  html.is-mhws-paged .mhws-print-head .sheet,
  html.is-mhws-paged .mhws-print-head .mhws-run-header .sheet {
    overflow: visible !important;
    min-height: 0 !important;
    height: auto !important;
    max-height: none !important;
    border: none !important;
    box-shadow: none !important;
  }
  .foot, footer.foot, html.pad-a4-full .foot {
    position: static !important;
    bottom: auto !important;
    width: 100% !important;
    margin: 0 !important;
    flex: none !important;
  }
  @media print {
    html.is-mhws-paged,
    html.is-mhws-paged.pad-a4-full,
    html.is-mhws-paged.mhws-compose-multipage {
      height: auto !important;
      max-height: none !important;
      min-height: 0 !important;
      overflow: visible !important;
    }
    html.is-mhws-paged body,
    html.is-mhws-paged.pad-a4-full body {
      height: auto !important;
      max-height: none !important;
      min-height: 0 !important;
      overflow: visible !important;
      background: #fff !important;
    }
    html.is-mhws-paged .mhws-run-header,
    html.is-mhws-paged .mhws-run-footer,
    html.is-mhws-paged .body-area,
    html.is-mhws-paged .body,
    html.is-mhws-paged body > img.wm,
    html.is-mhws-paged .mhws-run-header .sheet,
    html.is-mhws-paged .body-area .sheet {
      display: none !important;
      height: 0 !important;
      max-height: 0 !important;
      overflow: hidden !important;
      page-break-before: avoid !important;
      page-break-after: avoid !important;
    }
    html.is-mhws-paged .mhws-print-desk {
      padding: 0 !important;
      background: #fff !important;
    }
    html.is-mhws-paged .mhws-print-sheet {
      margin: 0 !important;
      box-shadow: none !important;
      page-break-after: avoid !important;
      break-after: avoid-page !important;
    }
    html.is-mhws-paged .mhws-print-sheet:not(:last-child) {
      page-break-after: always !important;
      break-after: page !important;
    }
  }
  html.pad-a4-full .slogan-bar { margin-left: -12mm !important; margin-right: -12mm !important; width: auto !important; }
  html.pad-a4-blank .slogan-bar { position: static !important; margin-left: -14mm !important; margin-right: -14mm !important; width: auto !important; }
</style>
"""

PRINT_CHROME_MM = {
    "none": (0.0, 0.0),
    "simple": (32.0, 22.0),
    "tpl-mhws-letterhead": (58.0, 32.0),
    "tpl-rwa-letterhead-blank": (48.0, 28.0),
}
PLAIN_CHROME_IDS = frozenset({"none", "none-template", "plain"})
DEFAULT_PRINT_CHROME_MM = (48.0, 28.0)


def compose_page_css(chrome_id: str, margins: dict[str, float], *, watermark: bool = True) -> str:
    mt, mr, mb, ml = margins["top"], margins["right"], margins["bottom"], margins["left"]
    wm_off = (
        ""
        if watermark
        else "img.wm { display: none !important; visibility: hidden !important; opacity: 0 !important; }"
    )
    return f"""
<style id="mhws-compose-page">
  @page {{
    size: 210mm 297mm;
    margin: 0;
  }}
  .body-area, .body {{
    padding: {mt:g}mm {mr:g}mm {mb:g}mm {ml:g}mm !important;
  }}
  img.wm {{
    position: fixed !important;
    top: 50% !important;
    left: 50% !important;
    transform: translate(-50%, -50%) !important;
    width: min(112mm, 70%) !important;
    max-height: 46% !important;
    height: auto !important;
    object-fit: contain !important;
    opacity: 0.75 !important;
    z-index: 0 !important;
    pointer-events: none !important;
    display: block !important;
  }}
  {wm_off}
  html.is-mhws-paged .mhws-run-header,
  html.is-mhws-paged .mhws-run-footer,
  html.is-mhws-paged .body-area,
  html.is-mhws-paged .body {{
    display: none !important;
  }}
  html.is-mhws-paged .mhws-run-header img.wm,
  html.is-mhws-paged body > img.wm,
  html.is-mhws-paged .sheet > img.wm {{
    display: none !important;
  }}
  html.is-mhws-paged .mhws-print-head .sheet,
  html.is-mhws-paged .mhws-print-head .mhws-run-header .sheet,
  .mhws-print-head .sheet {{
    overflow: visible !important;
    min-height: 0 !important;
    height: auto !important;
    max-height: none !important;
    width: 210mm !important;
    display: block !important;
    flex: none !important;
    border: none !important;
    box-shadow: none !important;
  }}
  .mhws-print-desk {{
    display: none;
    background: #c5cdd8;
    padding: 8mm 0;
  }}
  html.is-mhws-paged .mhws-print-desk {{ display: block; }}
  html.is-mhws-paged, html.is-mhws-paged body {{
    background: #c5cdd8 !important;
    width: auto !important;
    min-width: 100% !important;
  }}
  html.is-mhws-paged .mhws-print-desk {{
    width: 100%;
    box-sizing: border-box;
  }}
  .mhws-print-sheet {{
    position: relative;
    box-sizing: border-box;
    width: 210mm;
    min-height: 297mm;
    height: 297mm;
    margin: 0 auto 8mm;
    background: #fff;
    overflow: visible;
    display: flex;
    flex-direction: column;
    page-break-after: avoid;
    break-after: avoid-page;
    box-shadow: 0 1px 6px rgba(15, 40, 80, 0.2);
  }}
  .mhws-print-sheet:not(:last-child) {{
    page-break-after: always;
    break-after: page;
  }}
  .mhws-print-sheet:last-child {{ margin-bottom: 0; page-break-after: avoid; break-after: avoid-page; }}
  .mhws-print-head, .mhws-print-foot {{
    flex: 0 0 auto;
    width: 210mm;
    background: #fff;
    overflow: visible;
  }}
  .mhws-print-head .sheet {{
    border: none !important;
    box-shadow: none !important;
  }}
  .mhws-print-body {{
    flex: 1 1 auto;
    min-height: 0;
    overflow: hidden;
    padding: {mt:g}mm {mr:g}mm {mb:g}mm {ml:g}mm !important;
    box-sizing: border-box;
  }}
  .mhws-print-sheet img.wm {{
    position: absolute !important;
    left: 50% !important;
    top: 50% !important;
    transform: translate(-50%, -50%) !important;
    width: min(112mm, 70%) !important;
    max-height: 46% !important;
    height: auto !important;
    object-fit: contain !important;
    opacity: 0.75 !important;
    display: block !important;
    visibility: visible !important;
    pointer-events: none !important;
    z-index: 0 !important;
  }}
  .mhws-print-head,
  .mhws-print-body,
  .mhws-print-foot {{
    position: relative;
    z-index: 1;
  }}
  @media print {{
    html.is-mhws-paged,
    html.is-mhws-paged.pad-a4-full,
    html.is-mhws-paged.pad-a4-blank,
    html.is-mhws-paged.mhws-compose-multipage {{
      height: auto !important;
      max-height: none !important;
      min-height: 0 !important;
      overflow: visible !important;
    }}
    html.is-mhws-paged body,
    html.is-mhws-paged.pad-a4-full body,
    html.is-mhws-paged.mhws-compose-multipage body {{
      height: auto !important;
      max-height: none !important;
      min-height: 0 !important;
      width: auto !important;
      overflow: visible !important;
      margin: 0 !important;
      padding: 0 !important;
      background: #fff !important;
    }}
    html.is-mhws-paged .mhws-run-header,
    html.is-mhws-paged .mhws-run-footer,
    html.is-mhws-paged .body-area,
    html.is-mhws-paged .body,
    html.is-mhws-paged body > img.wm,
    html.is-mhws-paged .mhws-run-header .sheet,
    html.is-mhws-paged .body-area .sheet {{
      display: none !important;
      height: 0 !important;
      max-height: 0 !important;
      overflow: hidden !important;
      page-break-before: avoid !important;
      page-break-after: avoid !important;
      break-before: avoid-page !important;
      break-after: avoid-page !important;
    }}
    html.is-mhws-paged .mhws-print-desk {{
      padding: 0 !important;
      background: #fff !important;
    }}
    html.is-mhws-paged .mhws-print-sheet {{
      margin: 0 !important;
      box-shadow: none !important;
      height: 297mm !important;
      min-height: 297mm !important;
      max-height: 297mm !important;
      page-break-after: avoid !important;
      break-after: avoid-page !important;
    }}
    html.is-mhws-paged .mhws-print-sheet:not(:last-child) {{
      page-break-after: always !important;
      break-after: page !important;
    }}
  }}
</style>
"""


PREVIEW_PAGINATE_JS = r"""
<script id="mhws-preview-pager">
(function () {
  function go() {
    if (document.documentElement.classList.contains('is-mhws-paged')) return;
    var body = document.querySelector('.body-area, .body');
    var head = document.querySelector('.mhws-run-header');
    var foot = document.querySelector('.mhws-run-footer');
    if (!body) return;
    var wm = document.querySelector('img.wm');
    var probe = document.createElement('div');
    probe.style.cssText = 'position:absolute;left:-9999px;width:210mm;height:297mm;';
    document.body.appendChild(probe);
    var pagePx = probe.offsetHeight;
    probe.remove();
    var headH = head ? head.offsetHeight : 0;
    var footH = foot ? foot.offsetHeight : 0;
    var mt = parseFloat((body.style.paddingTop || getComputedStyle(body).paddingTop)) || 0;
    var mb = parseFloat((body.style.paddingBottom || getComputedStyle(body).paddingBottom)) || 0;
    var inner = Math.max(48, pagePx - headH - footH - mt - mb);
    var desk = document.createElement('div');
    desk.className = 'mhws-print-desk';
    var blocks = Array.prototype.slice.call(body.children);
    var slot = null;
    var used = 0;
    function newSheet() {
      var sheet = document.createElement('div');
      sheet.className = 'mhws-print-sheet';
      if (wm) {
        var w = wm.cloneNode(true);
        w.removeAttribute('id');
        sheet.appendChild(w);
      }
      if (head) {
        var h = head.cloneNode(true);
        h.className = 'mhws-print-head';
        h.style.position = 'static';
        sheet.appendChild(h);
      }
      var b = document.createElement('div');
      b.className = 'mhws-print-body';
      sheet.appendChild(b);
      if (foot) {
        var f = foot.cloneNode(true);
        f.className = 'mhws-print-foot';
        f.style.position = 'static';
        sheet.appendChild(f);
      }
      desk.appendChild(sheet);
      slot = b;
      used = 0;
    }
    newSheet();
    blocks.forEach(function (block) {
      var hgt = block.getBoundingClientRect().height;
      if (used > 8 && used + hgt > inner + 1) newSheet();
      slot.appendChild(block);
      used += hgt;
    });
    var sheets = desk.querySelectorAll('.mhws-print-sheet');
    for (var si = sheets.length - 1; si > 0; si -= 1) {
      var lb = sheets[si].querySelector('.mhws-print-body');
      if (lb && !lb.children.length) sheets[si].remove();
    }
    document.documentElement.classList.add('is-mhws-paged');
    document.body.appendChild(desk);
  }
  if (document.readyState === 'loading') document.addEventListener('DOMContentLoaded', go);
  else go();
})();
</script>
"""


def inject_compose_chrome_layout(html: str) -> str:
    text = html or ""
    if 'id="mhws-compose-chrome-layout"' in text:
        return text
    if "</head>" in text:
        return text.replace("</head>", COMPOSE_CHROME_LAYOUT_CSS + "\n</head>", 1)
    return text + COMPOSE_CHROME_LAYOUT_CSS


def inject_compose_page_extras(html: str, chrome_id: str, margins: dict[str, float], *, watermark: bool = True) -> str:
    text = html or ""
    css = compose_page_css(chrome_id, margins, watermark=watermark)
    if chrome_id not in PLAIN_CHROME_IDS:
        css += COMPOSE_CHROME_LAYOUT_CSS
    if 'id="mhws-compose-page"' not in text:
        if "</head>" in text:
            text = text.replace("</head>", css + "\n</head>", 1)
        else:
            text = css + text
    if 'id="mhws-preview-pager"' not in text:
        if "</body>" in text:
            text = text.replace("</body>", PREVIEW_PAGINATE_JS + "\n</body>", 1)
        else:
            text += PREVIEW_PAGINATE_JS
    return text


def inject_compose_pdf_css(html: str) -> str:
    if not html:
        return html
    if 'id="mhws-compose-pdf"' in html:
        return html
    if "</head>" in html:
        return html.replace("</head>", COMPOSE_PDF_CSS + "\n</head>", 1)
    return COMPOSE_PDF_CSS + html

_WM_IMG_RE = re.compile(r'<img\b[^>]*\bwm\b[^>]*>\s*', re.I)
_BODY_AREA_RE = re.compile(
    r'(<div\b[^>]*\bbody-area\b[^>]*>)(.*?)(</div>)',
    re.I | re.S,
)
_SCREEN_HINT_RE = re.compile(r'<p\s+class=["\']screen-hint["\'][\s\S]*?</p>', re.I)
_DATA_URI_RE = re.compile(
    r"^data:image/(png|jpe?g|gif|webp|svg\+xml);base64,(.+)$",
    re.I | re.S,
)


def as_bool(raw: Any, default: bool = True) -> bool:
    if raw is None:
        return default
    if isinstance(raw, bool):
        return raw
    if isinstance(raw, (int, float)):
        return bool(raw)
    return str(raw).strip().lower() not in {"", "0", "false", "no", "off", "none"}


def _find_body_area(html: str) -> tuple[int, str, str, str] | None:
    """Return (start, open_tag, inner_html, after_close) for the writing area."""
    match = re.search(r"<div\b[^>]*\bbody-area\b[^>]*>", html or "", flags=re.I)
    if not match:
        return None
    inner, after = _split_at_div_depth(html[match.end() :], 1)
    return match.start(), match.group(0), inner, after


def inject_body_area(pad_html: str, body: str) -> str | None:
    found = _find_body_area(pad_html or "")
    if not found:
        return None
    start, open_tag, _inner, after = found
    return pad_html[:start] + open_tag + body + "</div>" + after


def strip_print_pad_common(html: str) -> str:
    return re.sub(r'<link\b[^>]*print-pad-common[^>]*>\s*', "", html or "", flags=re.I)


def _add_html_class(html: str, class_name: str) -> str:
    def add(match: re.Match[str]) -> str:
        tag = match.group(0)
        if class_name in tag:
            return tag
        if re.search(r"\bclass\s*=", tag, re.I):
            return re.sub(r'(\bclass\s*=\s*["\'])', rf"\1{class_name} ", tag, count=1, flags=re.I)
        return tag[:-1] + f' class="{class_name}">'

    return re.sub(r"<html\b[^>]*>", add, html or "", count=1, flags=re.I)


def _split_at_div_depth(html: str, depth: int = 1) -> tuple[str, str]:
    current = depth
    for match in re.finditer(r"</?div\b[^>]*>", html or "", flags=re.I):
        token = match.group(0)
        closing = token.lower().startswith("</")
        self_close = token.rstrip().endswith("/>")
        if closing:
            current -= 1
            if current == 0:
                return html[: match.start()], html[match.end() :]
        elif not self_close:
            current += 1
    return html, ""


def _extract_wm(html: str) -> tuple[str, str]:
    found = _WM_IMG_RE.findall(html or "")
    return "".join(found), _WM_IMG_RE.sub("", html or "")


def _last_div_open(html: str, class_name: str) -> re.Match[str] | None:
    found_open = None
    for found in re.finditer(r"<div\b[^>]*>", html or "", flags=re.I):
        if re.search(rf"\b{re.escape(class_name)}\b", found.group(0), re.I):
            found_open = found
    return found_open


def paginate_pad_html(html: str) -> str:
    """Repeat letterhead header/footer on every printed page; let the body flow."""
    text = strip_print_pad_common(html or "")
    found = _find_body_area(text)
    if not found:
        return _add_html_class(text, "mhws-compose-multipage")
    start, _open_tag, body, suffix = found
    prefix = text[:start]
    pad_open = _last_div_open(prefix, "pad")
    sheet_open = _last_div_open(prefix, "sheet")
    container = pad_open or sheet_open
    if not container:
        return _add_html_class(text, "mhws-compose-multipage")
    header_inner = prefix[container.end() :]
    footer_inner, rest = _split_at_div_depth(suffix, 1)
    wm, header_inner = _extract_wm(header_inner)
    if pad_open and sheet_open:
        pre_pad = prefix[sheet_open.end() : pad_open.start()]
        extra_wm, pre_pad = _extract_wm(pre_pad)
        wm = extra_wm + wm
        header_inner = f"{pre_pad}<div class=\"pad\">{header_inner}</div>"
        footer_inner = f"<div class=\"pad\">{footer_inner}</div>"
        before = prefix[: sheet_open.end()]
    else:
        before = prefix[: container.end()]
    table = (
        f"{wm}"
        f'<div class="mhws-run-header">{header_inner}</div>'
        f'<div class="body-area">{body}</div>'
        f'<div class="mhws-run-footer">{footer_inner}</div>'
    )
    return _add_html_class(before + table + rest, "mhws-compose-multipage")


def _header_has_end_separator(html: str) -> bool:
    text = html or ""
    if re.search(r'<div class="rule"', text, flags=re.I):
        return True
    return "mhws-header-gold-rule" in text or "officers-foot" in text


def _extract_div_by_class(html: str, class_name: str) -> str:
    text = html or ""
    match = re.search(
        rf'<div\b[^>]*\bclass="[^"]*\b{re.escape(class_name)}\b[^"]*"[^>]*>',
        text,
        flags=re.I,
    )
    if not match:
        match = re.search(
            rf"<div\b[^>]*\bclass='[^']*\b{re.escape(class_name)}\b[^']*'[^>]*>",
            text,
            flags=re.I,
        )
    if not match:
        return ""
    inner, _rest = _split_at_div_depth(text[match.end() :], 1)
    return inner


def finalize_compose_pdf_html(html: str, margins: dict[str, float]) -> str:
    """Build a fixed A4 sheet for PDF export (mirrors preview pager without a browser)."""
    text = html or ""
    if "mhws-print-desk" in text or not re.search(r"\bmhws-run-header\b", text, flags=re.I):
        return text
    wm_match = re.search(r'<img\b[^>]*\bwm\b[^>]*>', text, flags=re.I)
    wm = wm_match.group(0) if wm_match else ""
    head = _extract_div_by_class(text, "mhws-run-header")
    body = _extract_div_by_class(text, "body-area")
    foot = _extract_div_by_class(text, "mhws-run-footer")
    desk = (
        '<div class="mhws-print-desk">'
        '<div class="mhws-print-sheet">'
        f"{wm}"
        f'<div class="mhws-print-head">{head}</div>'
        f'<div class="mhws-print-body">{body}</div>'
        f'<div class="mhws-print-foot">{foot}</div>'
        "</div></div>"
    )
    text = _add_html_class(text, "is-mhws-paged")
    if "</body>" in text:
        return text.replace("</body>", desk + "\n</body>", 1)
    return text + desk


def extract_pad_chrome(html: str) -> dict[str, str]:
    """Header/footer markup from a letterhead pad, for the composer page frames."""
    text = strip_print_pad_common(strip_screen_chrome(html or ""))
    css_parts = [m.group(1) for m in re.finditer(r"<style\b[^>]*>([\s\S]*?)</style>", text, flags=re.I)]
    chrome_css = "\n".join(css_parts)
    found = _find_body_area(text)
    if not found:
        return {"headerHtml": "", "footerHtml": "", "chromeCss": chrome_css, "watermarkUrl": ""}
    start, _open_tag, _body, suffix = found
    prefix = text[:start]
    pad_open = _last_div_open(prefix, "pad")
    sheet_open = _last_div_open(prefix, "sheet")
    container = pad_open or sheet_open
    if not container:
        return {"headerHtml": "", "footerHtml": "", "chromeCss": chrome_css, "watermarkUrl": ""}
    header_inner = prefix[container.end() :]
    footer_inner, _rest = _split_at_div_depth(suffix, 1)
    wm, header_inner = _extract_wm(header_inner)
    if pad_open and sheet_open:
        pre_pad = prefix[sheet_open.end() : pad_open.start()]
        extra_wm, pre_pad = _extract_wm(pre_pad)
        wm = extra_wm + wm
        sheet_tag = prefix[sheet_open.start() : sheet_open.end()]
        layout_m = re.search(r'data-layout=["\']([^"\']+)["\']', sheet_tag, flags=re.I)
        layout = (layout_m.group(1) if layout_m else "top").strip() or "top"
        header_inner = (
            f'{pre_pad}<div class="sheet" data-layout="{layout}">'
            f'<div class="pad">{header_inner}</div></div>'
        )
        footer_inner = f'<div class="pad">{footer_inner}</div>'
    wm_src = ""
    src = re.search(r'\bsrc=["\']([^"\']+)["\']', wm, flags=re.I)
    if src:
        wm_src = src.group(1)
    header_out = header_inner.strip()
    if re.search(r'\bclass="[^"]*\bofficers\b', header_out, flags=re.I):
        if not re.search(r'\bclass="[^"]*\bsheet\b', header_out, flags=re.I):
            header_out = f'<div class="sheet" data-layout="top"><div class="pad">{header_out}</div></div>'
    if header_out and not _header_has_end_separator(header_out):
        header_out = re.sub(
            r"(</div>\s*</div>\s*)$",
            rf"{DIAMOND_RULE_HTML}\1",
            header_out,
            count=1,
        )
        if not re.search(r'<div class="rule"', header_out, flags=re.I):
            header_out = f'{header_out}{DIAMOND_RULE_HTML}'
    return {
        "headerHtml": header_out,
        "footerHtml": footer_inner.strip(),
        "chromeCss": chrome_css,
        "watermarkUrl": wm_src,
    }


def strip_pager_markup(html: str) -> str:
    text = html or ""
    pattern = re.compile(
        r'<div\b[^>]*\b(mhws-page-spacer|mhws-page-chrome|mhws-page-frames)\b[^>]*>',
        re.I,
    )
    while True:
        match = pattern.search(text)
        if not match:
            return text
        _inner, after = _split_at_div_depth(text[match.end() :], 1)
        text = text[: match.start()] + after


def strip_screen_chrome(html: str) -> str:
    return _SCREEN_HINT_RE.sub("", html or "", count=1)


def rewrite_pad_urls(html: str) -> str:
    text = html or ""
    text = text.replace("../assets/", "/assets/")
    text = re.sub(r"""(\s(?:src|href)=["'])assets/""", r"\1/assets/", text, flags=re.I)
    return text


def set_html_title(html: str, title: str) -> str:
    heading = _html_escape((title or "Document").strip() or "Document")
    if re.search(r"<title\b", html or "", re.I):
        return re.sub(r"<title>[^<]*</title>", f"<title>{heading}</title>", html, count=1, flags=re.I)
    if "</head>" in (html or ""):
        return html.replace("</head>", f"<title>{heading}</title>\n</head>", 1)
    return html


def html_fragment_to_text(body_html: str) -> str:
    from rwa_templates import sanitize_compose_html

    html = sanitize_compose_html(body_html)
    html = re.sub(r"(?i)<br\s*/?>", "\n", html)
    html = re.sub(r"(?i)</(p|h[1-6]|li|blockquote|tr|div)>", "\n", html)
    html = re.sub(r"(?i)</t[dh]>", "\t", html)
    html = re.sub(r"(?i)<hr\s*/?>", "\n---\n", html)
    html = re.sub(r"(?i)<img\b[^>]*>", "[image]\n", html)
    html = re.sub(r"<[^>]+>", "", html)
    html = unescape(html).replace("\xa0", " ").replace("&nbsp;", " ")
    lines = [re.sub(r"[ \t]+", " ", line).rstrip() for line in html.splitlines()]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"


def _parse_width_pct(style: str, default: int = 40) -> int:
    match = re.search(r"width\s*:\s*(\d+(?:\.\d+)?)\s*%", style or "", re.I)
    if not match:
        return default
    try:
        pct = int(round(float(match.group(1))))
    except ValueError:
        return default
    return max(10, min(100, pct))


def _image_bytes_from_src(src: str, site_root: pathlib.Path | None) -> tuple[bytes, str] | None:
    raw = (src or "").strip()
    if not raw:
        return None
    raw = raw.split("?", 1)[0].split("#", 1)[0]
    data = _DATA_URI_RE.match(raw)
    if data:
        kind = data.group(1).lower()
        try:
            blob = base64.b64decode(data.group(2), validate=False)
        except Exception:
            return None
        ext = "jpg" if kind.startswith("jp") else ("png" if "png" in kind else ("gif" if "gif" in kind else "png"))
        return blob, ext
    if site_root is None:
        return None
    cleaned = raw.replace("\\", "/")
    while cleaned.startswith("../"):
        cleaned = cleaned[3:]
    cleaned = cleaned.lstrip("./")
    if cleaned.startswith("/"):
        cleaned = cleaned.lstrip("/")
    if not cleaned.startswith("assets/") and not cleaned.startswith("documents/"):
        if raw.startswith("/") or raw.startswith("../"):
            pass
        else:
            return None
    path = (pathlib.Path(site_root) / cleaned).resolve()
    root = pathlib.Path(site_root).resolve()
    if str(path).startswith(str(root)) and path.is_file():
        return path.read_bytes(), path.suffix.lstrip(".").lower() or "png"
    return None


_SRC_ATTR_RE = re.compile(r"""(\ssrc\s*=\s*)(['"])([^'"]+)\2""", re.I)
_BASE_TAG_RE = re.compile(r"<base\b[^>]*/?>", re.I)
_ASSET_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "svg": "image/svg+xml",
}


def strip_base_tag(html: str) -> str:
    return _BASE_TAG_RE.sub("", html or "", count=1)


def embed_local_asset_urls(
    html: str,
    site_root: pathlib.Path | None,
    *,
    max_bytes: int = 1_800_000,
) -> str:
    """Inline local /assets images so blob preview and PDF both see the logo."""
    if not html or site_root is None:
        return html

    def repl(match: re.Match[str]) -> str:
        prefix, quote, src = match.group(1), match.group(2), match.group(3)
        if (src or "").startswith(("data:", "http:", "https:", "blob:")):
            return match.group(0)
        got = _image_bytes_from_src(src, site_root)
        if not got:
            return match.group(0)
        blob, ext = got
        if len(blob) > max_bytes:
            return match.group(0)
        mime = _ASSET_MIME.get(ext, "image/png")
        return f"{prefix}{quote}data:{mime};base64,{base64.b64encode(blob).decode('ascii')}{quote}"

    return _SRC_ATTR_RE.sub(repl, html)


def _picture_stream(blob: bytes, ext: str) -> io.BytesIO | None:
    try:
        from PIL import Image  # type: ignore
    except Exception:
        return io.BytesIO(blob) if ext in {"png", "jpg", "jpeg", "gif"} else None
    try:
        im = Image.open(io.BytesIO(blob))
        if im.mode not in {"RGB", "L"}:
            im = im.convert("RGB")
        out = io.BytesIO()
        im.save(out, format="PNG")
        out.seek(0)
        return out
    except Exception:
        return None


class _DocxBuilder(HTMLParser):
    def __init__(self, document, *, usable_mm: float, site_root: pathlib.Path | None):
        super().__init__(convert_charrefs=True)
        self.doc = document
        self.usable_mm = usable_mm
        self.site_root = site_root
        self.skip_stack: list[str] = []
        self.bold = 0
        self.italic = 0
        self.underline = 0
        self.para = None
        self.in_table = 0
        self.rows: list[list[dict[str, str]]] = []
        self.cur_row: list[dict[str, str]] | None = None
        self.td_bits: list[str] | None = None
        self.td_style = ""
        self.img_width_pct = 40
        self.class_stack: list[str] = []

    def _classes(self, raw: str) -> set[str]:
        return {c for c in (raw or "").split() if c}

    def _has_class(self, name: str) -> bool:
        return any(name in item.split() for item in self.class_stack)

    def _ensure_para(self):
        if self.in_table:
            return
        if self.para is None:
            self.para = self.doc.add_paragraph()

    def _add_run(self, text: str):
        if self.td_bits is not None:
            self.td_bits.append(text)
            return
        self._ensure_para()
        run = self.para.add_run(text)
        run.bold = self.bold > 0 or self._has_class("title") or self._has_class("name")
        run.italic = self.italic > 0
        run.underline = self.underline > 0

    def _add_picture(self, src: str, width_pct: int | None = None, width_mm: float | None = None):
        from docx.shared import Mm

        got = _image_bytes_from_src(src, self.site_root)
        if not got:
            if self.td_bits is not None:
                self.td_bits.append("[image]")
            return
        blob, ext = got
        stream = _picture_stream(blob, ext)
        if stream is None:
            return
        if width_mm is None:
            pct = width_pct if width_pct is not None else self.img_width_pct
            width_mm = max(18.0, min(self.usable_mm, self.usable_mm * (pct / 100.0)))
        if self.in_table:
            return
        self._ensure_para()
        try:
            self.para.add_run().add_picture(stream, width=Mm(width_mm))
        except Exception:
            pass

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        tag = tag.lower()
        ad = {k.lower(): (v or "") for k, v in attrs}
        classes = self._classes(ad.get("class", ""))
        if self.skip_stack:
            return
        if tag in {"script", "style", "head", "svg", "noscript"} or classes & {
            "screen-hint",
            "layout-picker",
            "accent-edge",
            "accent-edge-thin",
        }:
            self.skip_stack.append(tag)
            return
        if tag not in {"img", "br", "hr", "meta", "link", "input", "col"}:
            self.class_stack.append(ad.get("class", ""))
        if tag in {"b", "strong"}:
            self.bold += 1
        elif tag in {"i", "em"}:
            self.italic += 1
        elif tag == "u":
            self.underline += 1
        elif tag == "br":
            if self.td_bits is not None:
                self.td_bits.append("\n")
            elif self.para is not None:
                self.para.add_run("\n")
        elif tag in {"p", "h1", "h2", "h3", "blockquote", "header", "footer"}:
            if not self.in_table:
                self.para = self.doc.add_paragraph()
                if tag in {"h1", "h2", "h3"}:
                    self.bold += 1
                if tag in {"h1", "h2", "header", "footer"} or classes & {"brand", "org", "foot"}:
                    try:
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
        elif tag == "li":
            if not self.in_table:
                self.para = self.doc.add_paragraph()
                self._add_run("• ")
        elif tag == "div" and classes & {"role", "contact"}:
            if not self.in_table:
                self.para = self.doc.add_paragraph()
        elif tag == "table":
            self.in_table += 1
            self.rows = []
            self.cur_row = None
            self.para = None
        elif tag == "tr":
            self.cur_row = []
        elif tag in {"td", "th"}:
            self.td_bits = []
            self.td_style = ad.get("style", "")
        elif tag == "span" and "mhws-img" in classes:
            self.img_width_pct = _parse_width_pct(ad.get("style", ""), default=int(ad.get("data-width") or 40) or 40)
        elif tag == "img":
            if "wm" in classes:
                return
            src = ad.get("src") or ""
            if "logo" in classes or "seal" in classes or self._has_class("org") or self._has_class("brand"):
                self._add_picture(src, width_mm=22.0)
                return
            width = self.img_width_pct
            style_w = _parse_width_pct(ad.get("style", ""), default=0)
            if style_w:
                width = style_w
            self._add_picture(src, width)

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if self.skip_stack:
            if self.skip_stack[-1] == tag:
                self.skip_stack.pop()
            return
        if tag in {"b", "strong"}:
            self.bold = max(0, self.bold - 1)
        elif tag in {"i", "em"}:
            self.italic = max(0, self.italic - 1)
        elif tag == "u":
            self.underline = max(0, self.underline - 1)
        elif tag in {"p", "h1", "h2", "h3", "blockquote", "li", "header", "footer"}:
            if tag in {"h1", "h2", "h3"}:
                self.bold = max(0, self.bold - 1)
            self.para = None
        elif tag == "span" and self._has_class("title"):
            self._add_run(" — ")
        elif tag in {"td", "th"}:
            text = unescape("".join(self.td_bits or [])).strip()
            if self.cur_row is not None:
                self.cur_row.append({"text": text, "style": self.td_style or ""})
            self.td_bits = None
            self.td_style = ""
        elif tag == "tr":
            if self.cur_row is not None:
                self.rows.append(self.cur_row)
            self.cur_row = None
        elif tag == "table":
            self.in_table = max(0, self.in_table - 1)
            rows = self.rows
            self.rows = []
            if rows:
                cols = max(len(r) for r in rows)
                table = self.doc.add_table(rows=len(rows), cols=max(1, cols))
                for i, row in enumerate(rows):
                    for j in range(cols):
                        spec = row[j] if j < len(row) else {"text": "", "style": ""}
                        if isinstance(spec, str):
                            spec = {"text": spec, "style": ""}
                        _apply_docx_cell(table.rows[i].cells[j], spec)
            self.para = None
        elif tag == "span":
            self.img_width_pct = 40
        if self.class_stack:
            self.class_stack.pop()

    def handle_data(self, data: str):
        if self.skip_stack:
            return
        text = data.replace("\xa0", " ")
        if not text or (not text.strip() and text != " "):
            if text == " ":
                self._add_run(" ")
            return
        self._add_run(text)


def wrapped_html_to_docx_bytes(*, html: str, site_root: pathlib.Path | None = None) -> bytes:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Mm
    except Exception as exc:
        raise ValueError("Word export needs python-docx on the server.") from exc

    from rwa_templates import parse_doc_margins_mm

    margins = parse_doc_margins_mm(html)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    left = max(12.0, margins["left"] * 0.6)
    right = max(12.0, margins["right"] * 0.6)
    section.top_margin = Mm(max(10.0, margins["top"] * 0.6))
    section.right_margin = Mm(right)
    section.bottom_margin = Mm(max(10.0, margins["bottom"] * 0.6))
    section.left_margin = Mm(left)
    usable = 210.0 - left - right

    builder = _DocxBuilder(doc, usable_mm=usable, site_root=site_root)
    builder.feed(html or "")
    builder.close()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def html_fragment_to_docx_bytes(
    *,
    title: str,
    body_html: str,
    site_root: pathlib.Path | None = None,
) -> bytes:
    """Back-compat: wrap a fragment in a minimal heading then convert."""
    heading = _html_escape((title or "").strip())
    html = f"<html><body><h1>{heading}</h1>{body_html or ''}</body></html>"
    return wrapped_html_to_docx_bytes(html=html, site_root=site_root)


def export_filename(title: str, suffix: str) -> str:
    from rwa_templates import _safe_attach_name

    return _safe_attach_name(title or "document", suffix)

